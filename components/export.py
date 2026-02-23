"""내보내기 기능 - Markdown, Word(.docx), JSON."""

from __future__ import annotations

import io
from datetime import datetime

import streamlit as st

from src.paper_state import get_paper_state


def _build_markdown() -> str:
    """현재 상태의 논문을 Markdown 문자열로 변환한다."""
    ps = get_paper_state()

    if ps.final_paper:
        return ps.final_paper

    parts: list[str] = []
    parts.append(f"# {ps.topic}\n")
    if ps.research_question:
        parts.append(f"**연구 질문**: {ps.research_question}\n")

    if ps.overview:
        parts.append(f"\n## 개요\n\n{ps.overview}\n")

    for sec in ps.sections:
        title = sec.title if isinstance(sec, object) and hasattr(sec, "title") else sec.get("title", "")
        parts.append(f"\n## {title}\n")
        desc = sec.description if hasattr(sec, "description") else sec.get("description", "")
        if desc:
            parts.append(f"*{desc}*\n")

        content = ps.draft_sections.get(title, "")
        if content:
            parts.append(f"\n{content}\n")

        subs = sec.subsections if hasattr(sec, "subsections") else sec.get("subsections", [])
        for sub in subs:
            sub_title = sub.get("title", "") if isinstance(sub, dict) else sub
            parts.append(f"\n### {sub_title}\n")

    return "\n".join(parts)


def _build_docx_bytes() -> bytes:
    """Markdown 내용을 Word 문서(.docx)로 변환한다."""
    from docx import Document
    from docx.shared import Pt

    md = _build_markdown()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Malgun Gothic"

    for line in md.split("\n"):
        stripped = line.strip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        elif stripped:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_export_buttons() -> None:
    """사이드바에 내보내기 버튼들을 렌더링한다."""
    st.subheader("내보내기")
    ps = get_paper_state()

    if not ps.topic:
        st.caption("주제를 설정하면 내보내기가 가능합니다.")
        return

    now = datetime.now().strftime("%Y%m%d_%H%M")
    safe_topic = ps.topic[:20].replace(" ", "_")

    md_content = _build_markdown()
    st.download_button(
        label="Markdown (.md)",
        data=md_content,
        file_name=f"review_{safe_topic}_{now}.md",
        mime="text/markdown",
        use_container_width=True,
    )

    try:
        docx_bytes = _build_docx_bytes()
        st.download_button(
            label="Word (.docx)",
            data=docx_bytes,
            file_name=f"review_{safe_topic}_{now}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    except ImportError:
        st.caption("Word 내보내기: `pip install python-docx`")
